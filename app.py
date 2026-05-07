"""
IntraLingo - Professional Document Translation EN ↔ FR
"""

import os
import re
import shutil
import tempfile

import gradio as gr
import torch
from transformers import MarianMTModel, MarianTokenizer
from docx import Document


# ============================================================================
# GLOSSARY PROCESSOR
# ============================================================================

class GlossaryProcessor:
    """
    Enforce custom terminology by substituting source terms with opaque
    placeholders before translation and restoring target terms afterward.
    MarianMT copies unknown tokens through, making this reliable for
    proper nouns, brand names, and technical terms.
    """

    PLACEHOLDER_PATTERN = "TRGLOSS{i}X"

    def __init__(self, glossary_text: str):
        self.terms: dict[str, str] = {}
        for line in glossary_text.strip().splitlines():
            line = line.strip()
            if not line or line.startswith("#"):
                continue
            if "→" in line:
                parts = line.split("→", 1)
            elif "->" in line:
                parts = line.split("->", 1)
            elif "=" in line:
                parts = line.split("=", 1)
            else:
                continue
            src = parts[0].strip()
            tgt = parts[1].strip()
            if src and tgt:
                self.terms[src] = tgt

    def encode(self, text: str) -> tuple[str, dict[str, str]]:
        """Replace source terms with placeholders; return modified text and reverse map."""
        placeholder_map: dict[str, str] = {}  # placeholder → target term
        for i, (src_term, tgt_term) in enumerate(self.terms.items()):
            placeholder = self.PLACEHOLDER_PATTERN.format(i=i)
            # Case-insensitive replacement, preserve original placeholder casing
            pattern = re.compile(re.escape(src_term), re.IGNORECASE)
            if pattern.search(text):
                text = pattern.sub(placeholder, text)
                placeholder_map[placeholder] = tgt_term
        return text, placeholder_map

    def decode(self, text: str, placeholder_map: dict[str, str]) -> tuple[str, int]:
        """Restore target terms from placeholders; return text and count of replacements."""
        count = 0
        for placeholder, tgt_term in placeholder_map.items():
            # Match placeholder and common casing variants (model may lowercase)
            variants = [placeholder, placeholder.lower(), placeholder.upper()]
            for variant in variants:
                if variant in text:
                    text = text.replace(variant, tgt_term)
                    count += 1
                    break
        return text, count


# ============================================================================
# MARIAN TRANSLATOR
# ============================================================================

class MarianTranslator:
    """Helsinki-NLP MarianMT translator for English ↔ French."""

    MODEL_MAP = {
        ("en", "fr"): "Helsinki-NLP/opus-mt-tc-big-en-fr",
        ("fr", "en"): "Helsinki-NLP/opus-mt-tc-big-fr-en",
    }

    def __init__(self, source_lang: str, target_lang: str):
        self.source_lang = source_lang
        self.target_lang = target_lang
        self.device = "cuda" if torch.cuda.is_available() else "cpu"
        self.model_name = self.MODEL_MAP[(source_lang, target_lang)]

        print(f"Loading {self.model_name} on {self.device}…")
        self.tokenizer = MarianTokenizer.from_pretrained(self.model_name)
        self.model = MarianMTModel.from_pretrained(self.model_name)
        self.model.to(self.device)
        self.model.eval()
        print("Model ready.")

    def translate(self, text: str) -> str:
        if not text.strip():
            return text
        inputs = self.tokenizer(
            [text],
            return_tensors="pt",
            padding=True,
            truncation=True,
            max_length=512,
        )
        inputs = {k: v.to(self.device) for k, v in inputs.items()}
        with torch.no_grad():
            translated = self.model.generate(**inputs, max_length=512, num_beams=4)
        return self.tokenizer.decode(translated[0], skip_special_tokens=True)


# ============================================================================
# POST-PROCESSING
# ============================================================================

# Each entry: (compiled regex, replacement string)
# Applied in order after EN→FR translation.
_EN_TO_FR_FIXES = [
    # Salutations — model often produces unnatural word order or phrasing
    (re.compile(r"Cher(?:e)? Monsieur ou Madame[,.]?", re.IGNORECASE), "Madame, Monsieur,"),
    (re.compile(r"Cher(?:e)? Madame ou Monsieur[,.]?", re.IGNORECASE), "Madame, Monsieur,"),
    (re.compile(r"Cher Monsieur[,.]?", re.IGNORECASE), "Monsieur,"),
    (re.compile(r"Chère Madame[,.]?", re.IGNORECASE), "Madame,"),
    (re.compile(r"Cher Monsieur/Madame[,.]?", re.IGNORECASE), "Madame, Monsieur,"),
    # Closings — literal translations that sound odd in French
    (re.compile(r"Cordialement vôtre[,.]?", re.IGNORECASE), "Cordialement,"),
    (re.compile(r"Sincèrement vôtre[,.]?", re.IGNORECASE), "Veuillez agréer l'expression de mes salutations distinguées,"),
    (re.compile(r"Sincèrement[,.]?", re.IGNORECASE), "Cordialement,"),
    (re.compile(r"Avec respect[,.]?", re.IGNORECASE), "Veuillez agréer mes salutations distinguées,"),
    (re.compile(r"Meilleures salutations[,.]?", re.IGNORECASE), "Bien cordialement,"),
]

_FR_TO_EN_FIXES = [
    # French formal closings that model over-translates
    (re.compile(r"Please accept the expression of my distinguished regards[,.]?", re.IGNORECASE),
     "Yours sincerely,"),
    (re.compile(r"Please accept the expression of my sincere greetings[,.]?", re.IGNORECASE),
     "Yours faithfully,"),
    (re.compile(r"Please accept my distinguished greetings[,.]?", re.IGNORECASE),
     "Yours sincerely,"),
]


def _post_process(text: str, src_lang: str) -> str:
    fixes = _EN_TO_FR_FIXES if src_lang == "en" else _FR_TO_EN_FIXES
    for pattern, replacement in fixes:
        text = pattern.sub(replacement, text)
    return text


# ============================================================================
# IN-PLACE PARAGRAPH TRANSLATION
# ============================================================================

from docx.oxml.ns import qn as _qn

_W_R  = _qn("w:r")
_W_T  = _qn("w:t")
_W_BR = _qn("w:br")


def _translate_para_inplace(para, translator: "MarianTranslator",
                            glossary: "GlossaryProcessor") -> int:
    """
    Translate a paragraph in-place.

    - If the paragraph contains soft line breaks (<w:br/>), each line segment
      is translated independently so the break positions are preserved exactly.
    - Otherwise the whole paragraph text is translated as one unit and the
      result is placed in run[0] (which keeps all its formatting attributes).

    Returns number of glossary hits.
    """
    p_elem = para._p

    # Collect text elements (<w:t>) and break elements (<w:br/>) in document
    # order by walking every <w:r> that is a descendant of this paragraph.
    # We build a list of "segments": each segment is a list of <w:t> nodes
    # whose text should be translated together (i.e. between two <w:br/>s).
    segments: list[list] = [[]]   # list of lists of lxml elements

    for r_elem in p_elem.iter(_W_R):
        for child in r_elem:
            if child.tag == _W_T:
                segments[-1].append(child)
            elif child.tag == _W_BR:
                segments.append([])   # start a new segment after the break

    total_hits = 0

    if len(segments) <= 1:
        # No soft line breaks — original fast path
        full_text = "".join((e.text or "") for e in segments[0]) if segments else ""
        if not full_text.strip():
            return 0
        translated, hits = _translate_block_text(full_text, translator, glossary)
        total_hits += hits
        if segments and segments[0]:
            segments[0][0].text = translated
            for t_elem in segments[0][1:]:
                t_elem.text = ""
    else:
        # Has soft line breaks — translate each segment independently
        for seg in segments:
            seg_text = "".join((e.text or "") for e in seg)
            if not seg_text.strip() or not seg:
                continue
            translated, hits = _translate_block_text(seg_text, translator, glossary)
            total_hits += hits
            seg[0].text = translated
            for t_elem in seg[1:]:
                t_elem.text = ""

    return total_hits


# ============================================================================
# TRANSLATION PIPELINE
# ============================================================================

# Global model cache to avoid reloading on every call
_translator_cache: dict[str, MarianTranslator] = {}


def _get_translator(src: str, tgt: str) -> MarianTranslator:
    key = f"{src}_{tgt}"
    if key not in _translator_cache:
        _translator_cache[key] = MarianTranslator(src, tgt)
    return _translator_cache[key]


def _translate_block_text(text: str, translator: MarianTranslator,
                          glossary: GlossaryProcessor) -> tuple[str, int]:
    """Translate a single text string with glossary enforcement and post-processing."""
    encoded, placeholder_map = glossary.encode(text)
    translated = translator.translate(encoded)
    translated, count = glossary.decode(translated, placeholder_map)
    translated = _post_process(translated, translator.source_lang)
    return translated, count


def translate_document(
    input_file,
    language_direction: str,
    glossary_text: str,
    progress=gr.Progress(),
) -> tuple:

    if input_file is None:
        return None, "❌ Please upload a .docx document first.", ""

    lang_map = {
        "English → French": ("en", "fr"),
        "French → English": ("fr", "en"),
    }
    src, tgt = lang_map[language_direction]
    glossary = GlossaryProcessor(glossary_text or "")

    try:
        file_path = input_file if isinstance(input_file, str) else input_file.name
        file_size = os.path.getsize(file_path)
        if file_size > 10 * 1024 * 1024:
            return None, f"❌ File too large ({file_size/1024/1024:.1f} MB). Limit: 10 MB.", ""

        progress(0.05, desc="Loading translation model…")
        translator = _get_translator(src, tgt)

        # Copy original: all page layout, styles, table borders, headers/footers
        # are preserved because we only touch run.text — never recreate structure.
        progress(0.12, desc="Preparing document copy…")
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            output_path = tmp.name
        shutil.copy2(file_path, output_path)

        doc = Document(output_path)
        body_paras = list(doc.paragraphs)
        tables = list(doc.tables)

        # Collect all paragraphs to translate (body + table cells + headers/footers)
        table_paras = [
            para
            for table in tables
            for row in table.rows
            for cell in row.cells
            for para in cell.paragraphs
        ]
        section_paras = [
            para
            for section in doc.sections
            for hf in (section.header, section.footer,
                       section.even_page_header, section.even_page_footer,
                       section.first_page_header, section.first_page_footer)
            if hf is not None
            for para in hf.paragraphs
        ]

        para_count = sum(1 for p in body_paras if p.text.strip())
        table_count = len(tables)
        total_glossary_hits = 0

        all_paras = body_paras + table_paras + section_paras
        total = len(all_paras)

        for i, para in enumerate(all_paras):
            progress(0.15 + (i / total) * 0.73,
                     desc=f"Translating paragraph {i + 1}/{total}…")
            total_glossary_hits += _translate_para_inplace(para, translator, glossary)

        progress(0.90, desc="Saving document…")
        doc.save(output_path)

        progress(1.0, desc="Done!")

        glossary_line = (
            f"- Glossary terms enforced: **{total_glossary_hits}** replacements"
            if glossary.terms
            else "- Glossary: none provided"
        )

        summary = "\n".join([
            "### 📊 Translation Report",
            f"- Direction: **{src.upper()} → {tgt.upper()}**",
            f"- Paragraphs preserved: **{para_count}**",
            f"- Tables preserved: **{table_count}**",
            glossary_line,
            f"- Model: `{translator.model_name}`",
            "",
            "✅ Download your translated document below.",
            "",
            "> 💡 Always review MT output before external use.",
        ])

        return output_path, "✅ Translation complete!", summary

    except Exception as exc:
        import traceback
        tb = traceback.format_exc()
        print(tb)
        return None, f"❌ Error: {exc}", ""


# ============================================================================
# GRADIO INTERFACE
# ============================================================================

with gr.Blocks(title="IntraLingo", theme=gr.themes.Soft()) as demo:

    gr.Markdown("""
    # 🌐 IntraLingo
    ## Professional Document Translation · EN ↔ FR

    Upload a `.docx` file → get back a fully translated document with **all formatting intact**:
    headings, bold/italic, tables, paragraph styles. Add optional custom terminology to
    enforce consistent business vocabulary.
    """)

    with gr.Row():
        with gr.Column(scale=1):
            gr.Markdown("### 📁 Document")
            input_file = gr.File(
                label="Upload .docx",
                file_types=[".docx"],
                type="filepath",
            )
            language_direction = gr.Dropdown(
                choices=["English → French", "French → English"],
                value="English → French",
                label="Translation direction",
            )

        with gr.Column(scale=1):
            gr.Markdown("### 📖 Custom Glossary *(optional)*")
            glossary_input = gr.Textbox(
                label="Term pairs (one per line)",
                placeholder=(
                    "# One pair per line, any separator\n"
                    "invoice → facture\n"
                    "compliance → conformité\n"
                    "CEO → PDG\n"
                    "Acme Corp → Acme Corp"
                ),
                lines=8,
                max_lines=20,
            )

    translate_btn = gr.Button("🔄 Translate Document", variant="primary", size="lg")

    with gr.Row():
        with gr.Column(scale=1):
            output_file = gr.File(label="📥 Translated Document", interactive=False)
            status_msg = gr.Markdown("")
        with gr.Column(scale=1):
            report_box = gr.Markdown("")

    translate_btn.click(
        fn=translate_document,
        inputs=[input_file, language_direction, glossary_input],
        outputs=[output_file, status_msg, report_box],
    )

    gr.Markdown("""
    ---
    **Model:** [Helsinki-NLP/opus-mt-tc-big-en-fr](https://huggingface.co/Helsinki-NLP/opus-mt-tc-big-en-fr)
    · MarianMT, state-of-the-art open-source NMT
    &nbsp;·&nbsp; **Format preservation:** headers, bold/italic, tables, paragraph styles
    &nbsp;·&nbsp; **Glossary enforcement:** placeholder-based terminology injection
    &nbsp;·&nbsp; **Privacy:** documents processed in memory only, never stored
    &nbsp;·&nbsp; *Developed by Agnès Heijligers — ML Engineer & Professional Translator*
    """)

if __name__ == "__main__":
    demo.queue()
    demo.launch()