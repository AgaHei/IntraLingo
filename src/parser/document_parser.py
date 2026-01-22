"""
IntraLingo - Document Parser Module
Handles parsing .docx files, preserving formatting, and reconstructing documents
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy


class DocumentParser:
    """
    Parses Word documents while preserving formatting for translation.
    Handles paragraphs, tables, lists, and text formatting (bold, italic, etc.)
    """
    
    def __init__(self, input_path):
        """
        Initialize parser with input document path.
        
        Args:
            input_path: Path to the .docx file to parse
        """
        self.input_path = input_path
        self.doc = Document(input_path)
        self.parsed_content = []
        
    def parse(self):
        """
        Parse the document into a structured format that preserves all formatting.
        Returns a list of content blocks (paragraphs and tables).
        """
        self.parsed_content = []
        
        # We need to track both paragraphs and tables in document order
        # python-docx doesn't provide direct access to mixed content order,
        # so we'll use the document body's elements
        
        # Track table elements to avoid parsing their internal paragraphs
        table_elements = set()
        for table in self.doc.tables:
            # Get all paragraph elements inside this table
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        table_elements.add(para._element)
        
        for element in self.doc.element.body:
            # Check if it's a paragraph
            if element.tag.endswith('p'):
                # Skip if this paragraph is inside a table
                if element not in table_elements:
                    para = self._get_paragraph_from_element(element)
                    if para is not None:
                        para_data = self._parse_paragraph(para)
                        self.parsed_content.append(para_data)
            
            # Check if it's a table
            elif element.tag.endswith('tbl'):
                table = self._get_table_from_element(element)
                if table is not None:
                    table_data = self._parse_table(table)
                    self.parsed_content.append(table_data)
        
        return self.parsed_content
    
    def _get_paragraph_from_element(self, element):
        """Get paragraph object from XML element."""
        for para in self.doc.paragraphs:
            if para._element == element:
                return para
        return None
    
    def _get_table_from_element(self, element):
        """Get table object from XML element."""
        for table in self.doc.tables:
            if table._element == element:
                return table
        return None
    
    def _parse_paragraph(self, para):
        """
        Parse a paragraph, preserving style and run-level formatting.
        
        Returns:
            dict: {
                'type': 'paragraph',
                'style': style name,
                'alignment': alignment,
                'spacing_before': spacing before,
                'spacing_after': spacing after,
                'line_spacing': line spacing,
                'runs': list of run dictionaries
            }
        """
        runs_data = []
        
        for run in para.runs:
            # Capture more font properties
            run_data = {
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_size': run.font.size,
                'font_name': run.font.name
            }
            
            # Capture font color if available
            try:
                if run.font.color.rgb:
                    run_data['font_color'] = run.font.color.rgb
            except:
                pass
                
            runs_data.append(run_data)

        # Capture paragraph spacing
        spacing_before = None
        spacing_after = None
        line_spacing = None
        
        try:
            if para.paragraph_format.space_before:
                spacing_before = para.paragraph_format.space_before
            if para.paragraph_format.space_after:
                spacing_after = para.paragraph_format.space_after
            if para.paragraph_format.line_spacing:
                line_spacing = para.paragraph_format.line_spacing
        except:
            pass

        return {
            'type': 'paragraph',
            'style': para.style.name,
            'alignment': para.alignment,
            'spacing_before': spacing_before,
            'spacing_after': spacing_after, 
            'line_spacing': line_spacing,
            'runs': runs_data
        }
    
    def _parse_table(self, table):
        """
        Parse a table, preserving cell content and formatting.
        
        Returns:
            dict: {
                'type': 'table',
                'rows': list of rows, each containing cells
            }
        """
        rows_data = []
        
        for row in table.rows:
            cells_data = []
            for cell in row.cells:
                # Each cell can contain multiple paragraphs
                cell_paragraphs = []
                for para in cell.paragraphs:
                    para_data = self._parse_paragraph(para)
                    cell_paragraphs.append(para_data)
                
                cells_data.append({
                    'paragraphs': cell_paragraphs
                })
            
            rows_data.append(cells_data)
        
        return {
            'type': 'table',
            'rows': rows_data,
            'num_rows': len(table.rows),
            'num_cols': len(table.columns)
        }
    
    def reconstruct(self, translated_content, output_path):
        """
        Reconstruct a Word document from translated content,
        preserving all original formatting.
        
        Args:
            translated_content: List of content blocks (same structure as parsed_content)
            output_path: Path where to save the reconstructed document
        """
        # Create new document
        new_doc = Document()
        
        # Copy styles from original document
        self._copy_styles(new_doc)
        
        for i, block in enumerate(translated_content):
            if block['type'] == 'paragraph':
                self._reconstruct_paragraph(new_doc, block)
            elif block['type'] == 'table':
                self._reconstruct_table(new_doc, block)
        
        new_doc.save(output_path)
        return output_path
    
    def _copy_styles(self, new_doc):
        """Copy styles from original document to new document."""
        # Copy built-in styles and custom styles
        try:
            for style in self.doc.styles:
                if style.name not in [s.name for s in new_doc.styles]:
                    try:
                        # Create new style with same properties
                        if style.type == 1:  # Paragraph style
                            new_style = new_doc.styles.add_style(style.name, 1)
                            if hasattr(style, 'font'):
                                new_style.font.name = style.font.name
                                new_style.font.size = style.font.size
                                new_style.font.bold = style.font.bold
                                new_style.font.italic = style.font.italic
                    except:
                        # Skip if style creation fails
                        continue
        except:
            # If style copying fails, continue with defaults
            pass
    
    def _reconstruct_paragraph(self, doc, para_data):
        """
        Reconstruct a paragraph with all its formatting.
        
        Args:
            doc: Document object to add paragraph to
            para_data: Parsed paragraph data dictionary
        """
        para = doc.add_paragraph()
        
        # Try to set style, but handle errors gracefully
        try:
            para.style = para_data['style']
        except Exception:
            # Use default style if there's an error
            pass
        
        # Set alignment
        if para_data['alignment'] is not None:
            para.alignment = para_data['alignment']
            
        # Set paragraph spacing
        try:
            if para_data.get('spacing_before'):
                para.paragraph_format.space_before = para_data['spacing_before']
            if para_data.get('spacing_after'):
                para.paragraph_format.space_after = para_data['spacing_after'] 
            if para_data.get('line_spacing'):
                para.paragraph_format.line_spacing = para_data['line_spacing']
        except:
            pass
        
        # Add runs with formatting
        for run_data in para_data['runs']:
            run = para.add_run(run_data['text'])
            
            # Apply formatting safely
            if run_data['bold'] is not None:
                try:
                    run.bold = run_data['bold']
                except:
                    pass
            if run_data['italic'] is not None:
                try:
                    run.italic = run_data['italic']
                except:
                    pass
            if run_data['underline'] is not None:
                try:
                    run.underline = run_data['underline']
                except:
                    pass
                    
            # Apply font name
            if run_data.get('font_name'):
                try:
                    run.font.name = run_data['font_name']
                except:
                    pass
                    
            # Apply font color
            if run_data.get('font_color'):
                try:
                    run.font.color.rgb = run_data['font_color']
                except:
                    pass
            
            # Handle font size properly
            try:
                from docx.shared import Pt
                
                if run_data['font_size'] is not None:
                    original_size = run_data['font_size']
                    
                    # Handle different font size formats
                    if hasattr(original_size, 'pt'):
                        # Already a Pt object
                        run.font.size = original_size
                    elif isinstance(original_size, (int, float)):
                        # Convert to reasonable point size
                        if original_size > 1000:  # Likely in twips
                            size_in_points = original_size / 20
                        elif original_size > 100:  # Likely in half-points or other unit
                            size_in_points = original_size / 2
                        else:  # Already in points
                            size_in_points = original_size
                        
                        # Ensure reasonable size (6-72 points)
                        size_in_points = max(6, min(72, size_in_points))
                        run.font.size = Pt(size_in_points)
                    else:
                        # Use style defaults
                        self._apply_style_font_size(run, para_data['style'])
                else:
                    # Font size is None - use style-appropriate defaults
                    self._apply_style_font_size(run, para_data['style'])
            except Exception as e:
                # Fallback to default
                try:
                    run.font.size = Pt(11)
                except:
                    pass
                    
    def _apply_style_font_size(self, run, style_name):
        """Apply appropriate font size based on style."""
        try:
            from docx.shared import Pt
            if style_name in ['Heading 1', 'Title']:
                run.font.size = Pt(16)
            elif style_name in ['Heading 2']:
                run.font.size = Pt(14)
            elif style_name in ['Heading 3']:
                run.font.size = Pt(13)
            else:
                run.font.size = Pt(11)  # Standard document size
        except:
            pass
    
    def _reconstruct_table(self, doc, table_data):
        """
        Reconstruct a table with all its content and formatting.
        
        Args:
            doc: Document object to add table to
            table_data: Parsed table data dictionary
        """
        # Create table with correct dimensions
        table = doc.add_table(rows=table_data['num_rows'], 
                             cols=table_data['num_cols'])
        
        # Apply basic table styling
        table.style = 'Table Grid'  # Ensures visible borders
        
        # Fill table cells
        for i, row_data in enumerate(table_data['rows']):
            for j, cell_data in enumerate(row_data):
                cell = table.rows[i].cells[j]
                
                # Clear default paragraph
                cell.paragraphs[0].clear()
                
                # Add all paragraphs from original cell
                for k, para_data in enumerate(cell_data['paragraphs']):
                    if k == 0:
                        # Use existing first paragraph
                        para = cell.paragraphs[0]
                        try:
                            para.style = para_data['style']
                        except:
                            pass
                        if para_data['alignment'] is not None:
                            para.alignment = para_data['alignment']
                            
                        # Apply paragraph spacing
                        try:
                            if para_data.get('spacing_before'):
                                para.paragraph_format.space_before = para_data['spacing_before']
                            if para_data.get('spacing_after'):
                                para.paragraph_format.space_after = para_data['spacing_after'] 
                        except:
                            pass
                    else:
                        # Add new paragraph
                        para = cell.add_paragraph()
                        try:
                            para.style = para_data['style']
                        except:
                            pass
                        if para_data['alignment'] is not None:
                            para.alignment = para_data['alignment']
                    
                    # Add runs with full formatting
                    for run_data in para_data['runs']:
                        run = para.add_run(run_data['text'])
                        
                        # Apply all formatting properties
                        if run_data['bold'] is not None:
                            try:
                                run.bold = run_data['bold']
                            except:
                                pass
                        if run_data['italic'] is not None:
                            try:
                                run.italic = run_data['italic']
                            except:
                                pass
                        if run_data['underline'] is not None:
                            try:
                                run.underline = run_data['underline']
                            except:
                                pass
                        if run_data.get('font_name'):
                            try:
                                run.font.name = run_data['font_name']
                            except:
                                pass
                        if run_data.get('font_color'):
                            try:
                                run.font.color.rgb = run_data['font_color']
                            except:
                                pass
                        
                        # Handle font size
                        try:
                            from docx.shared import Pt
                            if run_data['font_size'] is not None:
                                if hasattr(run_data['font_size'], 'pt'):
                                    run.font.size = run_data['font_size']
                                else:
                                    # Convert to appropriate size
                                    size = run_data['font_size']
                                    if isinstance(size, (int, float)) and 6 <= size <= 72:
                                        run.font.size = Pt(size)
                                    elif size > 1000:  # Likely twips
                                        run.font.size = Pt(size / 20)
                        except:
                            pass


def dummy_translate_text(text):
    """
    Dummy translation function for testing.
    Adds [TRANSLATED] prefix to text.
    
    Args:
        text: Text to "translate"
    
    Returns:
        str: "Translated" text
    """
    if text.strip():
        return f"[TRANSLATED] {text}"
    return text


def translate_content(parsed_content, translation_func=dummy_translate_text):
    """
    Apply translation function to all text in parsed content.
    
    Args:
        parsed_content: List of parsed content blocks
        translation_func: Function that takes text and returns translated text
    
    Returns:
        List of translated content blocks with same structure
    """
    translated = []
    
    for block in parsed_content:
        if block['type'] == 'paragraph':
            translated_block = copy.deepcopy(block)
            for run in translated_block['runs']:
                run['text'] = translation_func(run['text'])
            translated.append(translated_block)
        
        elif block['type'] == 'table':
            translated_block = copy.deepcopy(block)
            for row in translated_block['rows']:
                for cell in row:
                    for para in cell['paragraphs']:
                        for run in para['runs']:
                            run['text'] = translation_func(run['text'])
            translated.append(translated_block)
    
    return translated
